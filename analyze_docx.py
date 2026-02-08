#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import json

docx_path = r'c:\Users\Adam\Documents\Ktransport\Ktransporte\Présentation Ktransport.docx'

try:
    doc = Document(docx_path)

    structure = {
        'total_paragraphs': len(doc.paragraphs),
        'paragraphs': []
    }

    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        structure['paragraphs'].append({
            'index': i,
            'text': text,
            'is_empty': not text,
            'length': len(text)
        })

    # Save structure to JSON file
    with open(r'c:\Users\Adam\Documents\Ktransport\Ktransporte\document_structure.json', 'w', encoding='utf-8') as f:
        json.dump(structure, f, ensure_ascii=False, indent=2)

    # Also create a readable text file
    with open(r'c:\Users\Adam\Documents\Ktransport\Ktransporte\document_structure.txt', 'w', encoding='utf-8') as f:
        f.write("=== DOCUMENT STRUCTURE ===\n\n")
        f.write(f"Total paragraphs: {len(doc.paragraphs)}\n\n")
        for item in structure['paragraphs']:
            if item['is_empty']:
                f.write(f"[{item['index']}] (EMPTY PARAGRAPH)\n")
            else:
                preview = item['text'][:80] + \
                    "..." if len(item['text']) > 80 else item['text']
                f.write(f"[{item['index']}] {preview}\n")
            f.write("\n")

    print(f"Analysis complete. Found {len(doc.paragraphs)} paragraphs.")
    print(f"Saved to: document_structure.json and document_structure.txt")

except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
