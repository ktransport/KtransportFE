#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
import sys

docx_path = r'c:\Users\Adam\Documents\Ktransport\Ktransporte\Présentation Ktransport.docx'

try:
    doc = Document(docx_path)

    # Create mapping file
    output = []
    output.append("=== CURRENT DOCUMENT STRUCTURE ===\n")
    output.append(f"Total paragraphs in document: {len(doc.paragraphs)}\n\n")

    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        if text:
            preview = text[:100] + "..." if len(text) > 100 else text
            output.append(f"Position {i}: {preview}\n")
        else:
            output.append(f"Position {i}: [EMPTY]\n")

    # Write to file
    with open(r'c:\Users\Adam\Documents\Ktransport\Ktransporte\current_structure.txt', 'w', encoding='utf-8') as f:
        f.writelines(output)

    # Also print to console
    print("".join(output))
    print(f"\nStructure saved to: current_structure.txt")

except Exception as e:
    error_msg = f"Error: {e}\n"
    print(error_msg)
    import traceback
    traceback.print_exc()
    with open(r'c:\Users\Adam\Documents\Ktransport\Ktransporte\error_log.txt', 'w', encoding='utf-8') as f:
        f.write(error_msg)
        f.write(traceback.format_exc())
